import os
import torch
import transformers
from transformers import AutoModelForCausalLM, AutoTokenizer
import pdfplumber
import docx
import html2text
import re

class AdvancedFileConverter:
    def __init__(self, file_path):
        """
        Initialize the converter with a file path and LLM model
        
        Args:
            file_path (str): Path to the file to be converted
        """
        print(f"[DEBUG] Initializing AdvancedFileConverter for file: {file_path}")
        self.file_path = file_path
        
        # Initialize the LLM model and pipeline
        print("[DEBUG] Loading Llama 3 model and tokenizer...")
        self.model = AutoModelForCausalLM.from_pretrained("llama3_3b_instruct")
        self.tokenizer = AutoTokenizer.from_pretrained("llama3_3b_instruct")
        print("[DEBUG] Initializing text generation pipeline...")
        self.pipeline = transformers.pipeline(
                            "text-generation",
                            model="meta-llama/Llama-3.2-3B-Instruct",
                            model_kwargs={"torch_dtype": torch.bfloat16},
                        )
        print("[DEBUG] Model and pipeline initialized successfully")

    def convert_to_markdown(self):
        """
        Convert file to markdown based on file extension
        
        Returns:
            str: Markdown-formatted content
        """
        file_extension = os.path.splitext(self.file_path)[1].lower()
        print(f"[DEBUG] Detected file extension: {file_extension}")

        # Extract raw text based on file type
        try:
            if file_extension == '.pdf':
                print("[DEBUG] Extracting text from PDF...")
                raw_text = self._extract_pdf_text()
            elif file_extension == '.docx':
                print("[DEBUG] Extracting text from DOCX...")
                raw_text = self._extract_docx_text()
            elif file_extension == '.html':
                print("[DEBUG] Extracting text from HTML...")
                raw_text = self._extract_html_text()
            elif file_extension in ['.txt', '.md']:
                print("[DEBUG] Reading text directly from file...")
                with open(self.file_path, 'r', encoding='utf-8') as file:
                    raw_text = file.read()
            else:
                raise ValueError(f"Unsupported file type: {file_extension}")
            
            print(f"[DEBUG] Raw text extracted. Length: {len(raw_text)} characters")

            # Convert to markdown using LLM
            print("[DEBUG] Beginning markdown conversion with LLM...")
            markdown_text = self._convert_to_markdown_with_llm(raw_text)
            
            print(f"[DEBUG] Markdown conversion complete. Length: {len(markdown_text)} characters")
            return markdown_text

        except Exception as e:
            print(f"[ERROR] Error during markdown conversion: {e}")
            raise

    def _extract_pdf_text(self):
        """
        Extract text from PDF while preserving original formatting
        
        Returns:
            str: Extracted text from PDF
        """
        text = ""
        with pdfplumber.open(self.file_path) as pdf:
            total_pages = len(pdf.pages)
            print(f"[DEBUG] PDF contains {total_pages} pages")
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                text += page_text + "\n\n"
                print(f"[DEBUG] Extracted text from page {page_num}")
        
        return text.strip()

    def _extract_docx_text(self):
        """
        Extract text from DOCX while preserving formatting
        
        Returns:
            str: Extracted text from DOCX
        """
        doc = docx.Document(self.file_path)
        text = ""
        paragraph_count = len(doc.paragraphs)
        print(f"[DEBUG] DOCX contains {paragraph_count} paragraphs")
        
        for para_num, para in enumerate(doc.paragraphs, 1):
            # Preserve original paragraph structure
            if para.style.name.startswith('Heading'):
                # Handle headings
                heading_level = int(para.style.name.replace('Heading', ''))
                text += f"{'#' * heading_level} {para.text}\n\n"
                print(f"[DEBUG] Found heading (Level {heading_level}): {para.text}")
            else:
                text += f"{para.text}\n\n"
            
            if para_num % 10 == 0:
                print(f"[DEBUG] Processed {para_num} paragraphs")
        
        return text.strip()

    def _extract_html_text(self):
        """
        Convert HTML to markdown using html2text
        
        Returns:
            str: Markdown-converted HTML
        """
        h = html2text.HTML2Text()
        h.ignore_links = False
        h.ignore_images = False
        
        with open(self.file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        print(f"[DEBUG] HTML file size: {len(html_content)} characters")
        markdown_text = h.handle(html_content).strip()
        print(f"[DEBUG] Converted HTML to markdown. Length: {len(markdown_text)} characters")
        
        return markdown_text

    def _convert_to_markdown_with_llm(self, raw_text):
        """
        Use LLM to format text into markdown while preserving content
        
        Args:
            raw_text (str): Original text to be formatted
        
        Returns:
            str: Markdown-formatted text
        """
        # Prompt engineered to convert text to markdown without content modification
        prompt = f"""Convert the following text to markdown format. 
        Rules:
        1. Do NOT modify or paraphrase any content
        2. Preserve all original text exactly
        3. Apply markdown formatting:
           - Use headers (#, ##, ###) for structure
           - Use bullet points (*) where appropriate
           - Add code block formatting (```) for code
           - Preserve line breaks
           - Use italics (*) and bold (**) for emphasis if detected

        Original Text:
        {raw_text}

        Markdown Version:"""

        print("[DEBUG] Generating markdown using LLM...")
        print(f"[DEBUG] Prompt length: {len(prompt)} characters")
        print(f"[DEBUG] Max new tokens: {len(raw_text) * 2}")

        # Generate markdown using the pipeline
        response = self.pipeline(
            prompt, 
            max_new_tokens=len(raw_text) * 2,  # Allow sufficient tokens
            do_sample=False,  # Minimize randomness
            temperature=0.1   # Low temperature for consistency
        )[0]['generated_text']

        print("[DEBUG] LLM response received")

        # Extract just the markdown part (after the prompt)
        markdown_text = response.split(prompt)[-1].strip()

        print(f"[DEBUG] Markdown text length: {len(markdown_text)} characters")

        # Save markdown to file for reference
        with open('markdownText.md', 'w', encoding='utf-8') as file:
            file.write(markdown_text)
        print("[DEBUG] Markdown text saved to markdownText.md")
        
        return markdown_text
    
class MarkdownChunker:
    def __init__(self, markdown_content, max_chunk_size=2000):
        """
        Initialize markdown chunker with content and optional max chunk size
        
        Args:
            markdown_content (str): Full markdown text
            max_chunk_size (int): Maximum size of each chunk
        """
        print(f"[DEBUG] Initializing MarkdownChunker. Content length: {len(markdown_content)}")
        print(f"[DEBUG] Max chunk size: {max_chunk_size}")
        self.markdown_content = markdown_content
        self.max_chunk_size = max_chunk_size

    def chunk_by_headings(self):
        """
        Chunk markdown content by headings
        
        Returns:
            list: Chunks of markdown text, each starting with a heading
        """
        print("[DEBUG] Beginning chunk_by_headings method")
        
        # Regular expression to find all headings (# to ######)
        heading_pattern = r'^(#{1,6})\s+(.+)$'
        
        # Split content into sections based on headings
        sections = []
        current_section = []
        current_heading_level = 0

        # Split the content into lines for precise processing
        lines = self.markdown_content.split('\n')
        print(f"[DEBUG] Total lines in markdown content: {len(lines)}")
        
        for line_num, line in enumerate(lines, 1):
            # Check if the line is a heading
            heading_match = re.match(heading_pattern, line, re.MULTILINE)
            
            if heading_match:
                # If we have a previous section, add it to sections
                if current_section:
                    sections.append('\n'.join(current_section))
                    print(f"[DEBUG] Completed section. Section length: {len(sections[-1])} characters")
                    current_section = []
                
                # Start a new section with this heading
                current_section.append(line)
                current_heading_level = len(heading_match.group(1))
                print(f"[DEBUG] Found heading (Level {current_heading_level}) at line {line_num}: {line}")
            else:
                # Add non-heading lines to the current section
                current_section.append(line)
        
        # Add the last section
        if current_section:
            sections.append('\n'.join(current_section))
            print(f"[DEBUG] Added final section. Length: {len(sections[-1])} characters")
        
        # Further chunk large sections if they exceed max_chunk_size
        final_chunks = []
        for section_num, section in enumerate(sections, 1):
            if len(section) <= self.max_chunk_size:
                final_chunks.append(section)
                print(f"[DEBUG] Section {section_num} fits within chunk size")
            else:
                # If a section is too large, split it while trying to preserve structure
                print(f"[DEBUG] Section {section_num} too large. Splitting...")
                split_chunks = self._split_large_section(section)
                final_chunks.extend(split_chunks)
                print(f"[DEBUG] Split into {len(split_chunks)} chunks")
        
        print(f"[DEBUG] Total chunks created: {len(final_chunks)}")
        return final_chunks

    def _split_large_section(self, section):
        """
        Split a large section into smaller chunks while preserving markdown structure
        
        Args:
            section (str): Large markdown section to split
        
        Returns:
            list: Smaller chunks of the section
        """
        print(f"[DEBUG] Splitting large section. Original length: {len(section)}")
        
        chunks = []
        current_chunk = []
        current_chunk_size = 0
        
        for line in section.split('\n'):
            line_length = len(line)
            
            # If adding this line would exceed max chunk size, start a new chunk
            if current_chunk_size + line_length > self.max_chunk_size:
                chunks.append('\n'.join(current_chunk))
                print(f"[DEBUG] Created chunk. Length: {len(chunks[-1])} characters")
                current_chunk = []
                current_chunk_size = 0
            
            current_chunk.append(line)
            current_chunk_size += line_length + 1  # +1 for newline
        
        # Add the last chunk if not empty
        if current_chunk:
            chunks.append('\n'.join(current_chunk))
            print(f"[DEBUG] Created final chunk. Length: {len(chunks[-1])} characters")
        
        return chunks

def read_file_with_encoding(file_path):
    """
    Attempt to read file with multiple encodings
    
    Args:
        file_path (str): Path to the file to read
    
    Returns:
        str: File content as a string
    """
    # List of encodings to try
    encodings = [
        'utf-8', 
        'latin-1',  # Most permissive encoding
        'cp1252',   # Windows encoding
        'iso-8859-1',
        'utf-16',
    ]
    
    print(f"[DEBUG] Attempting to read file: {file_path}")
    
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as file:
                content = file.read()
            print(f"[DEBUG] Successfully read file with {encoding} encoding")
            return content
        except UnicodeDecodeError:
            print(f"[DEBUG] Failed to read file with {encoding} encoding")
        except Exception as e:
            print(f"[DEBUG] Unexpected error with {encoding} encoding: {e}")
    
    # Fallback to binary read if all else fails
    try:
        with open(file_path, 'rb') as file:
            content = file.read().decode('latin-1')
        print("[DEBUG] Fell back to binary read with latin-1 encoding")
        return content
    except Exception as e:
        print(f"[ERROR] Could not read file: {e}")
        raise

def process_file_by_headings(file_path, max_chunk_size=2000):
    """
    Process a markdown file into chunks by headings
    
    Args:
        file_path (str): Path to the markdown file
        max_chunk_size (int): Maximum size of each chunk
    
    Returns:
        list: Chunks of markdown text
    """
    print(f"[DEBUG] Starting file processing: {file_path}")
    
    # Read the markdown content using robust reading method
    markdown_content = read_file_with_encoding(file_path)
    
    # Create converter to ensure markdown format
    converter = AdvancedFileConverter(file_path)
    markdown_content = converter.convert_to_markdown()
    
    # Create chunker and chunk by headings
    chunker = MarkdownChunker(markdown_content, max_chunk_size)
    chunks = chunker.chunk_by_headings()
    
    print(f"[DEBUG] Completed file processing. Total chunks: {len(chunks)}")
    return chunks

# Demonstration
if __name__ == "__main__":
    file_path = './data/onepage.pdf'  # Replace with your markdown file path
    markdown_chunks = process_file_by_headings(file_path)
    
    for idx, chunk in enumerate(markdown_chunks, 1):
        print(f"Chunk {idx} (Size: {len(chunk)} characters):\n{chunk}\n{'-' * 50}")